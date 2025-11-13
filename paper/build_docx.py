#!/usr/bin/env python3
"""Generate a single-column Word (.docx) document from IEEE_Paper_Document.md

Two generation modes:
1. If python-docx is available, use it for higher-fidelity formatting.
2. If python-docx is NOT available (PEP 668 / externally managed env), fall back to a pure builder
   that crafts the minimal DOCX (OpenXML) package manually without external deps.

Implemented content features:
- Headings (#, ##, ###) -> Heading1/Heading2/Heading3 styles (or manual style tags in fallback)
- Paragraphs
- Bullet / numbered lists (flattened; fallback just paragraphs prefixed with • / number)
- Code blocks (monospace style; fallback uses a custom Code style)
- Markdown pipe tables -> Word tables (fallback constructs <w:tbl>)

Limitations (both modes):
- Images, math, cross-references not rendered; left as plain text.
- Nested lists flattened.
- No automatic page breaks or sectioning beyond single section.

Output: IEEE_Paper.docx in the same folder.
"""
import re
import zipfile
import xml.sax.saxutils as saxutils
from pathlib import Path

try:
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore
    USE_DOCX = True
except Exception:  # python-docx not available
    USE_DOCX = False

ROOT = Path(__file__).parent
MARKDOWN = ROOT / 'IEEE_Paper_Document.md'
OUTPUT = ROOT / 'IEEE_Paper.docx'

HEADING_MAP = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}

code_block = False
code_lines = []

def add_code_block_docx(document, lines):
    """Add a code block using python-docx."""
    if not lines:
        return
    p = document.add_paragraph()
    run = p.add_run('\n'.join(lines))
    font = run.font
    font.name = 'Courier New'
    font.size = Pt(9)
    # Alignment left (default) - keep simple.

def escape_xml(text: str) -> str:
    return saxutils.escape(text)

def build_styles_xml():
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:style w:type="paragraph" w:styleId="Normal">'
        '<w:name w:val="Normal"/><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/></w:rPr>'
        '</w:style>'
        '<w:style w:type="paragraph" w:styleId="Heading1">'
        '<w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="32"/></w:rPr>'
        '</w:style>'
        '<w:style w:type="paragraph" w:styleId="Heading2">'
        '<w:name w:val="heading 2"/><w:basedOn w:val="Normal"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="28"/></w:rPr>'
        '</w:style>'
        '<w:style w:type="paragraph" w:styleId="Heading3">'
        '<w:name w:val="heading 3"/><w:basedOn w:val="Normal"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="24"/></w:rPr>'
        '</w:style>'
        '<w:style w:type="paragraph" w:styleId="Code">'
        '<w:name w:val="Code"/><w:basedOn w:val="Normal"/><w:rPr><w:rFonts w:ascii="Courier New" w:hAnsi="Courier New"/><w:sz w:val="18"/></w:rPr>'
        '</w:style>'
        '</w:styles>'
    )

def make_paragraph_xml(text: str, style: str | None = None):
    text = escape_xml(text)
    style_part = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ''
    return f'<w:p>{style_part}<w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'

def make_table_xml(rows: list[list[str]]):
    tbl_rows = []
    for r in rows:
        cells_xml = []
        for c in r:
            c_esc = escape_xml(c)
            cells_xml.append(
                '<w:tc><w:p><w:r><w:t xml:space="preserve">' + c_esc + '</w:t></w:r></w:p></w:tc>'
            )
        tbl_rows.append('<w:tr>' + ''.join(cells_xml) + '</w:tr>')
    return '<w:tbl>' + ''.join(tbl_rows) + '</w:tbl>'

def assemble_doc_xml(blocks: list[dict]):
    body_parts = []
    for b in blocks:
        kind = b['type']
        if kind == 'heading':
            body_parts.append(make_paragraph_xml(b['text'], b['style']))
        elif kind == 'paragraph':
            body_parts.append(make_paragraph_xml(b['text']))
        elif kind == 'code':
            body_parts.append(make_paragraph_xml('\n'.join(b['lines']), 'Code'))
        elif kind == 'table':
            body_parts.append(make_table_xml(b['rows']))
        elif kind == 'list_item':
            # Fallback list styling: prefix bullet/number
            prefix = b.get('prefix', '')
            body_parts.append(make_paragraph_xml(f'{prefix}{b["text"]}'))
    body = ''.join(body_parts) + '<w:sectPr/>'
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:body>{body}</w:body></w:document>'
    )

def write_minimal_docx(blocks: list[dict], output: Path):
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
        '</Types>'
    )
    rels_root = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        '</Relationships>'
    )
    doc_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>'
    )
    document_xml = assemble_doc_xml(blocks)
    styles_xml = build_styles_xml()
    with zipfile.ZipFile(output, 'w') as z:
        z.writestr('[Content_Types].xml', content_types)
        z.writestr('_rels/.rels', rels_root)
        z.writestr('word/document.xml', document_xml)
        z.writestr('word/_rels/document.xml.rels', doc_rels)
        z.writestr('word/styles.xml', styles_xml)

def add_code_block_fallback(blocks, lines):
    if not lines:
        return
    blocks.append({'type': 'code', 'lines': lines[:]})


def parse_table_block(block_lines):
    # First line header, second may be separator, rest data
    rows = []
    for raw in block_lines:
        line = raw.strip()
        if not line or set(line) <= {'|', '-', ':'}:
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    return rows


def is_table_line(line):
    return line.strip().startswith('|') and '|' in line.strip()[1:]


def flush_table_docx(document, table_lines):
    rows = parse_table_block(table_lines)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ''
            cell = table.cell(r_idx, c_idx)
            cell.text = text

def flush_table_fallback(blocks, table_lines):
    rows = parse_table_block(table_lines)
    if not rows:
        return
    blocks.append({'type': 'table', 'rows': rows})


def main():
    if not MARKDOWN.exists():
        raise SystemExit(f'Markdown source not found: {MARKDOWN}')
    with MARKDOWN.open('r', encoding='utf-8') as f:
        lines = f.readlines()

    # Unified parsing into intermediate blocks
    blocks: list[dict] = []
    table_buffer: list[str] = []
    global code_block, code_lines
    code_block = False
    code_lines = []

    for i, raw in enumerate(lines):
        line = raw.rstrip('\n')

        # Code blocks
        if line.strip().startswith('```'):
            if code_block:
                # close
                if USE_DOCX:
                    blocks.append({'type': 'code', 'lines': code_lines[:]})
                else:
                    add_code_block_fallback(blocks, code_lines)
                code_lines.clear()
                code_block = False
            else:
                code_block = True
            continue
        if code_block:
            code_lines.append(line)
            continue

        # Tables
        if is_table_line(line):
            table_buffer.append(line)
            next_line = lines[i+1] if i+1 < len(lines) else ''
            if not is_table_line(next_line):
                rows = parse_table_block(table_buffer)
                if rows:
                    blocks.append({'type': 'table', 'rows': rows})
                table_buffer.clear()
            continue

        # Blank line -> paragraph separator
        if not line.strip():
            blocks.append({'type': 'paragraph', 'text': ''})
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            style = HEADING_MAP.get(level, 'Heading 3') if USE_DOCX else f'Heading{min(level,3)}'
            blocks.append({'type': 'heading', 'text': text, 'style': style})
            continue

        # Bullet list
        if re.match(r'^\s*[-*]\s+\S', line):
            item = re.sub(r'^\s*[-*]\s+', '', line)
            if USE_DOCX:
                blocks.append({'type': 'list_item', 'text': item, 'list_kind': 'bullet'})
            else:
                blocks.append({'type': 'list_item', 'text': item, 'prefix': '• '})
            continue

        # Numbered list
        m_num = re.match(r'^\s*(\d+)\.\s+\S', line)
        if m_num:
            item = re.sub(r'^\s*\d+\.\s+', '', line)
            if USE_DOCX:
                blocks.append({'type': 'list_item', 'text': item, 'list_kind': 'number'})
            else:
                num = m_num.group(1)
                blocks.append({'type': 'list_item', 'text': item, 'prefix': f'{num}. '})
            continue

        # Default paragraph
        blocks.append({'type': 'paragraph', 'text': line})

    # If code block never closed
    if code_block and code_lines:
        if USE_DOCX:
            blocks.append({'type': 'code', 'lines': code_lines[:]})
        else:
            add_code_block_fallback(blocks, code_lines)

    if USE_DOCX:
        # Build with python-docx
        document = Document()
        # Normal style font
        try:
            normal = document.styles['Normal']
            normal.font.name = 'Times New Roman'
            normal.font.size = Pt(11)
        except Exception:
            pass

        for b in blocks:
            if b['type'] == 'heading':
                document.add_paragraph(b['text'], style=b['style'])
            elif b['type'] == 'paragraph':
                document.add_paragraph(b['text'])
            elif b['type'] == 'code':
                add_code_block_docx(document, b['lines'])
            elif b['type'] == 'table':
                flush_table_docx(document, ['|' + '|'.join(r) + '|' for r in b['rows']])
            elif b['type'] == 'list_item':
                p = document.add_paragraph(b['text'])
                kind = b.get('list_kind')
                if kind == 'bullet':
                    p.style = document.styles.get('List Bullet', document.styles['Normal'])
                elif kind == 'number':
                    p.style = document.styles.get('List Number', document.styles['Normal'])
        document.save(OUTPUT)
    else:
        # Fallback pure OpenXML
        write_minimal_docx(blocks, OUTPUT)

    print(f'Generated DOCX (mode={"python-docx" if USE_DOCX else "fallback"}): {OUTPUT}')

if __name__ == '__main__':
    main()
